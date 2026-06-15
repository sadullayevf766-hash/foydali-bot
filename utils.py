"""Foydali amallar: Rasm->PDF, PDF birlashtirish/bo'lish, rasm siqish/kichraytirish, QR, valyuta, buxgalter."""
import io
import os
import re
import html
import zipfile
from datetime import date, timedelta
import httpx
import qrcode
from PIL import Image
from pypdf import PdfReader, PdfWriter
from config import CBU_URL, BASE_DIR


def images_to_pdf(image_bytes_list: list[bytes]) -> bytes:
    """Bir nechta rasm baytlarini bitta PDF'ga aylantiradi."""
    images = []
    for b in image_bytes_list:
        img = Image.open(io.BytesIO(b))
        if img.mode != "RGB":
            img = img.convert("RGB")
        images.append(img)
    if not images:
        raise ValueError("Rasm yo'q")
    out = io.BytesIO()
    images[0].save(out, format="PDF", save_all=True, append_images=images[1:])
    return out.getvalue()


def merge_pdfs(pdf_bytes_list: list[bytes]) -> bytes:
    """Bir nechta PDF faylni bitta PDF'ga birlashtiradi."""
    writer = PdfWriter()
    for b in pdf_bytes_list:
        reader = PdfReader(io.BytesIO(b))
        for page in reader.pages:
            writer.add_page(page)
    out = io.BytesIO()
    writer.write(out)
    return out.getvalue()


def make_qr(text: str) -> bytes:
    """Matn yoki havoladan QR-kod rasmi (PNG) yasaydi."""
    qr = qrcode.QRCode(box_size=10, border=2)
    qr.add_data(text)
    qr.make(fit=True)
    img = qr.make_image(fill_color="black", back_color="white")
    out = io.BytesIO()
    img.save(out, format="PNG")
    return out.getvalue()


async def _fetch_latest_rates(client: httpx.AsyncClient) -> list:
    """CBU dan eng so'nggi e'lon qilingan kursni oladi (cbu.uz bosh sahifasi bilan bir xil).

    CBU kurslarni bir necha kun oldindan e'lon qiladi. Shu sababli bugundan
    bir necha kun keyingi sana bo'yicha so'rov yuboramiz — endpoint o'sha
    sanadan oldingi eng so'nggi mavjud kursni qaytaradi."""
    target = (date.today() + timedelta(days=7)).isoformat()
    url = f"https://cbu.uz/uz/arkhiv-kursov-valyut/json/all/{target}/"
    resp = await client.get(url)
    resp.raise_for_status()
    data = resp.json()
    if data:
        return data
    # Zaxira: oddiy (joriy) manzil
    resp = await client.get(CBU_URL)
    resp.raise_for_status()
    return resp.json()


async def get_rates(codes: list[str] | None = None) -> str:
    """Markaziy bank kurslarini matn ko'rinishida qaytaradi."""
    codes = codes or ["USD", "EUR", "RUB", "KZT"]
    async with httpx.AsyncClient(timeout=15) as client:
        data = await _fetch_latest_rates(client)
    by_code = {item["Ccy"]: item for item in data}
    lines = ["💱 *Markaziy bank rasmiy kursi*\n_(1 birlik = O'zbek so'mi)_\n"]
    for code in codes:
        item = by_code.get(code)
        if not item:
            continue
        rate = _fmt_som(item.get("Rate", "—"))
        diff = item.get("Diff") or "0"
        try:
            d = float(diff)
            arrow = "🔺" if d > 0 else ("🔻" if d < 0 else "▪️")
        except (ValueError, TypeError):
            arrow = "▪️"
        lines.append(f"{arrow} *{code}* — {rate} so'm  ({diff})")
    lines.append(f"\n📅 Sana: {data[0]['Date']}")
    lines.append("ℹ️ Bu — rasmiy kurs. Bank yoki shoxobchada narx biroz farq qilishi mumkin.")
    return "\n".join(lines)


def _fmt_som(rate_str: str) -> str:
    """Raqamni o'qishga qulay formatga keltiradi: 12014.48 -> '12 014'."""
    try:
        v = float(rate_str)
    except (ValueError, TypeError):
        return str(rate_str)
    if v >= 100:
        return f"{v:,.0f}".replace(",", " ")
    return f"{v:.2f}"


async def get_rate_map() -> dict:
    """Valyuta kodlari -> kurs (float) lug'atini qaytaradi (konvertor uchun)."""
    async with httpx.AsyncClient(timeout=15) as client:
        data = await _fetch_latest_rates(client)
    m = {}
    for it in data:
        try:
            m[it["Ccy"]] = float(it["Rate"])
        except (ValueError, TypeError, KeyError):
            continue
    return m, data[0].get("Date", "")


# ---------- Buxgalter: Son -> so'zda ----------

_UNITS = ["", "bir", "ikki", "uch", "to'rt", "besh", "olti", "yetti", "sakkiz", "to'qqiz"]
_TENS = ["", "o'n", "yigirma", "o'ttiz", "qirq", "ellik", "oltmish", "yetmish", "sakson", "to'qson"]
_SCALES = ["", "ming", "million", "milliard", "trillion", "kvadrillion"]


def _three_digit_words(x: int) -> str:
    """0..999 ni so'zga aylantiradi."""
    parts = []
    h, t, u = x // 100, (x % 100) // 10, x % 10
    if h:
        parts.append((_UNITS[h] + " " if h > 1 else "") + "yuz")
    if t:
        parts.append(_TENS[t])
    if u:
        parts.append(_UNITS[u])
    return " ".join(parts)


def number_to_uzbek_words(n: int) -> str:
    """Butun sonni o'zbekcha so'zlarga aylantiradi. Masalan 1250000 ->
    'bir million ikki yuz ellik ming'."""
    if n == 0:
        return "nol"
    if n < 0:
        return "minus " + number_to_uzbek_words(-n)
    groups = []
    while n > 0:
        groups.append(n % 1000)
        n //= 1000
    parts = []
    for i in range(len(groups) - 1, -1, -1):
        g = groups[i]
        if g == 0:
            continue
        if i == 1 and g == 1:
            parts.append("ming")  # "bir ming" emas, "ming"
        else:
            words = _three_digit_words(g)
            if i > 0:
                words += " " + _SCALES[i]
            parts.append(words)
    return " ".join(parts)


# ---------- Fayl/rasm vositalari ----------

def split_pdf(pdf_bytes: bytes):
    """PDF ni alohida sahifalarga bo'lib, ZIP qilib qaytaradi. (zip_bytes, sahifalar_soni)."""
    reader = PdfReader(io.BytesIO(pdf_bytes))
    buf = io.BytesIO()
    with zipfile.ZipFile(buf, "w", zipfile.ZIP_DEFLATED) as z:
        for i, page in enumerate(reader.pages, 1):
            writer = PdfWriter()
            writer.add_page(page)
            pb = io.BytesIO()
            writer.write(pb)
            z.writestr(f"sahifa_{i}.pdf", pb.getvalue())
    return buf.getvalue(), len(reader.pages)


def compress_image(image_bytes: bytes) -> bytes:
    """Rasm hajmini kamaytiradi (JPEG sifatini pasaytirib)."""
    img = Image.open(io.BytesIO(image_bytes))
    if img.mode in ("RGBA", "P", "LA"):
        img = img.convert("RGB")
    out = io.BytesIO()
    img.save(out, format="JPEG", quality=45, optimize=True)
    return out.getvalue()


def resize_image(image_bytes: bytes, max_side: int = 1024) -> bytes:
    """Rasmning eng uzun tomonini max_side ga keltiradi (proporsiya saqlanadi)."""
    img = Image.open(io.BytesIO(image_bytes))
    if img.mode in ("RGBA", "P", "LA"):
        img = img.convert("RGB")
    w, h = img.size
    scale = max_side / max(w, h)
    if scale < 1:
        img = img.resize((max(1, int(w * scale)), max(1, int(h * scale))))
    out = io.BytesIO()
    img.save(out, format="JPEG", quality=85, optimize=True)
    return out.getvalue()


# ---------- Matn -> Word / PDF hujjat ----------
#
# Tuzilma belgilari (talaba oddiy yozadi):
#   # Sarlavha          -> asosiy sarlavha (markazda, qalin, katta)
#   ## Bo'lim           -> bo'lim sarlavhasi (qalin)
#   ### Kichik bo'lim   -> kichik sarlavha (qalin, kursiv)
#   - punkt   yoki  * punkt   -> belgili ro'yxat
#   1. punkt                  -> raqamli ro'yxat
#   oddiy matn               -> xatboshi (paragraf)
#   bo'sh qator              -> ajratish

def _parse_blocks(text: str):
    """Matnni bloklarga ajratadi: (tur, matn). Tur: h1/h2/h3/bullet/number/body/blank."""
    blocks = []
    for raw in text.strip().split("\n"):
        s = raw.strip()
        if not s:
            blocks.append(("blank", ""))
        elif s.startswith("### "):
            blocks.append(("h3", s[4:].strip()))
        elif s.startswith("## "):
            blocks.append(("h2", s[3:].strip()))
        elif s.startswith("# "):
            blocks.append(("h1", s[2:].strip()))
        elif s[:2] in ("- ", "* ") or s.startswith("• "):
            blocks.append(("bullet", s[2:].strip()))
        elif re.match(r"^\d+[.)]\s+", s):
            blocks.append(("number", re.sub(r"^\d+[.)]\s+", "", s)))
        else:
            blocks.append(("body", s))
    # Agar foydalanuvchi '#' ishlatmagan bo'lsa, birinchi qisqa qatorni sarlavha qilamiz
    if not any(b[0] == "h1" for b in blocks):
        for i, (t, txt) in enumerate(blocks):
            if t == "blank":
                continue
            if t == "body" and len(txt) <= 80 and any(b[0] != "blank" for b in blocks[i + 1:]):
                blocks[i] = ("h1", txt)
            break
    return blocks


def text_to_docx(text: str) -> bytes:
    """Matndan tuzilmali Word (.docx) yasaydi: sarlavha, bo'limlar, ro'yxatlar, xatboshilar."""
    from docx import Document
    from docx.shared import Pt, Inches
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    doc = Document()
    normal = doc.styles["Normal"]
    normal.font.name = "Times New Roman"
    normal.font.size = Pt(14)

    def styled_run(p, txt, size, bold=False, italic=False):
        r = p.add_run(txt)
        r.bold = bold
        r.italic = italic
        r.font.size = Pt(size)
        r.font.name = "Times New Roman"

    for typ, txt in _parse_blocks(text):
        if typ == "blank":
            continue
        if typ == "h1":
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            styled_run(p, txt, 16, bold=True)
            p.paragraph_format.space_after = Pt(12)
        elif typ == "h2":
            p = doc.add_paragraph()
            styled_run(p, txt, 14, bold=True)
            p.paragraph_format.space_before = Pt(8)
            p.paragraph_format.space_after = Pt(4)
        elif typ == "h3":
            p = doc.add_paragraph()
            styled_run(p, txt, 14, bold=True, italic=True)
        elif typ == "bullet":
            doc.add_paragraph(txt, style="List Bullet")
        elif typ == "number":
            doc.add_paragraph(txt, style="List Number")
        else:  # body
            p = doc.add_paragraph(txt)
            p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            p.paragraph_format.line_spacing = 1.5
            p.paragraph_format.first_line_indent = Inches(0.5)

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


def text_to_pdf(text: str) -> bytes:
    """Matndan tuzilmali PDF yasaydi (o'zbek/kirill uchun DejaVu shrift)."""
    from reportlab.lib.pagesizes import A4
    from reportlab.lib.units import cm
    from reportlab.lib.styles import ParagraphStyle
    from reportlab.lib.enums import TA_JUSTIFY, TA_CENTER, TA_LEFT
    from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, ListFlowable, ListItem
    from reportlab.pdfbase import pdfmetrics
    from reportlab.pdfbase.ttfonts import TTFont

    if "DejaVu" not in pdfmetrics.getRegisteredFontNames():
        pdfmetrics.registerFont(TTFont("DejaVu", os.path.join(BASE_DIR, "fonts", "DejaVuSans.ttf")))

    body_style = ParagraphStyle("body", fontName="DejaVu", fontSize=14, leading=21,
                                alignment=TA_JUSTIFY, firstLineIndent=1 * cm, spaceAfter=6)
    h1_style = ParagraphStyle("h1", fontName="DejaVu", fontSize=16, leading=22,
                              alignment=TA_CENTER, spaceAfter=12)
    h2_style = ParagraphStyle("h2", fontName="DejaVu", fontSize=14, leading=20,
                              alignment=TA_LEFT, spaceBefore=8, spaceAfter=4)
    h3_style = ParagraphStyle("h3", fontName="DejaVu", fontSize=13, leading=18,
                              alignment=TA_LEFT, spaceBefore=4, spaceAfter=4)
    list_style = ParagraphStyle("li", fontName="DejaVu", fontSize=14, leading=20)

    def esc(t):
        return html.escape(t)

    story = []
    pending = []  # ro'yxat punktlarini to'plash

    def flush_list(ordered):
        nonlocal pending
        if pending:
            story.append(ListFlowable(
                [ListItem(Paragraph(esc(x), list_style)) for x in pending],
                bulletType="1" if ordered else "bullet", leftIndent=18))
            story.append(Spacer(1, 4))
            pending = []

    blocks = _parse_blocks(text)
    i = 0
    while i < len(blocks):
        typ, txt = blocks[i]
        if typ in ("bullet", "number"):
            ordered = (typ == "number")
            pending = []
            while i < len(blocks) and blocks[i][0] == typ:
                pending.append(blocks[i][1])
                i += 1
            flush_list(ordered)
            continue
        if typ == "blank":
            story.append(Spacer(1, 8))
        elif typ == "h1":
            story.append(Paragraph(f"<b>{esc(txt)}</b>", h1_style))
        elif typ == "h2":
            story.append(Paragraph(f"<b>{esc(txt)}</b>", h2_style))
        elif typ == "h3":
            story.append(Paragraph(f"<b><i>{esc(txt)}</i></b>", h3_style))
        else:
            story.append(Paragraph(esc(txt), body_style))
        i += 1

    buf = io.BytesIO()
    SimpleDocTemplate(buf, pagesize=A4, topMargin=2 * cm, bottomMargin=2 * cm,
                      leftMargin=2.5 * cm, rightMargin=1.5 * cm).build(story)
    return buf.getvalue()
